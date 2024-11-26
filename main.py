from docx import Document

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('Exam Date Sheet', level=1)

# Ask user for the number of exams
num_exams = int(input("Enter the number of exams: "))

# Add a table
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'

# Add the header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Date'
hdr_cells[1].text = 'Time'
hdr_cells[2].text = 'Course'
hdr_cells[3].text = 'Session'
hdr_cells[4].text = 'Room'
hdr_cells[5].text = 'Teacher Name'

# Collect data for each exam
for i in range(num_exams):
    print(f"\nEnter details for exam {i + 1}:")
    date = input("Enter the date (e.g., 25-Nov-2024): ")
    time = input("Enter the time (e.g., 10:00 AM): ")
    course = input("Enter the course name: ")
    session = input("Enter the semester & section (e.g., Fall 2024, Section A): ")
    room = input("Enter the room number: ")
    teacher_name = input("Enter the teacher's name: ")

    # Add a new row to the table
    row_cells = table.add_row().cells
    row_cells[0].text = date
    row_cells[1].text = time
    row_cells[2].text = course
    row_cells[3].text = session
    row_cells[4].text = room
    row_cells[5].text = teacher_name

# Save the document
doc_name = 'Date_Sheet.docx'
doc.save(doc_name)
print(f"\nDate sheet successfully created and saved as '{doc_name}'!")
